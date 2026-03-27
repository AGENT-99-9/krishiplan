#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════╗
║           SOIL REPORT EXTRACTOR — Single-File Edition              ║
║                                                                      ║
║  Extracts structured data from soil test reports using OCR.          ║
║  Supports: .png .jpg .jpeg .pdf .docx                               ║
║  Backends: pytesseract, easyocr (auto-detected)                     ║
║                                                                      ║
║  Usage:                                                              ║
║    python soil_extractor.py --check                                  ║
║    python soil_extractor.py --create-sample                          ║
║    python soil_extractor.py -i input/sample_soil_report.png          ║
║    python soil_extractor.py -i report.pdf -e easyocr                 ║
║    python soil_extractor.py -d ./input/ -w 4                         ║
╚══════════════════════════════════════════════════════════════════════╝
"""

# ═══════════════════════════════════════════════════════════════════════
#  IMPORTS
# ═══════════════════════════════════════════════════════════════════════
import argparse
import io
import json
import logging
import os
import platform
import re
import shutil
import subprocess
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

_MISSING = []
try:
    import numpy as np
except ImportError:
    _MISSING.append("numpy")
try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    _MISSING.append("Pillow")
try:
    import cv2
except ImportError:
    _MISSING.append("opencv-python")

if _MISSING:
    print(f"\n  Missing: {', '.join(_MISSING)}")
    print(f"  Run:  pip install {' '.join(_MISSING)}\n")
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════
#  CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════
BASE_DIR = Path(__file__).parent.resolve()
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"
LOG_DIR = BASE_DIR / "logs"

for _d in (INPUT_DIR, OUTPUT_DIR, LOG_DIR):
    _d.mkdir(parents=True, exist_ok=True)

SUPPORTED_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}
SUPPORTED_PDF_EXT = {".pdf"}
SUPPORTED_WORD_EXT = {".docx"}
SUPPORTED_EXT = SUPPORTED_IMAGE_EXT | SUPPORTED_PDF_EXT | SUPPORTED_WORD_EXT

TESSERACT_CONFIG = "--oem 3 --psm 6"
EASYOCR_LANGUAGES = ["en"]
OCR_DPI = 300
MAX_WORKERS = 4
MAX_FILE_SIZE_MB = 100
DENOISE_STRENGTH = 10
THRESHOLD_BLOCK_SIZE = 11
THRESHOLD_C = 2


# ═══════════════════════════════════════════════════════════════════════
#  WINDOWS TESSERACT FINDER
# ═══════════════════════════════════════════════════════════════════════
def _find_tesseract_windows() -> Optional[str]:
    """Search common Windows install locations for tesseract.exe."""
    if platform.system() != "Windows":
        return None

    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\ProgramData\chocolatey\bin\tesseract.exe",
        os.path.expanduser(
            r"~\AppData\Local\Tesseract-OCR\tesseract.exe"),
        os.path.expanduser(
            r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
        os.path.expanduser(
            r"~\scoop\apps\tesseract\current\tesseract.exe"),
        r"C:\Tesseract-OCR\tesseract.exe",
        r"D:\Tesseract-OCR\tesseract.exe",
    ]

    path_dirs = os.environ.get("PATH", "").split(os.pathsep)
    for d in path_dirs:
        p = os.path.join(d, "tesseract.exe")
        if p not in candidates:
            candidates.append(p)

    try:
        import winreg
        for hive in [winreg.HKEY_LOCAL_MACHINE,
                     winreg.HKEY_CURRENT_USER]:
            for subkey in [r"SOFTWARE\Tesseract-OCR",
                           r"SOFTWARE\WOW6432Node\Tesseract-OCR"]:
                try:
                    key = winreg.OpenKey(hive, subkey)
                    install_dir, _ = winreg.QueryValueEx(
                        key, "InstallDir")
                    winreg.CloseKey(key)
                    reg_path = os.path.join(
                        install_dir, "tesseract.exe")
                    if reg_path not in candidates:
                        candidates.insert(0, reg_path)
                except (FileNotFoundError, OSError):
                    pass
    except ImportError:
        pass

    try:
        flags = 0
        if hasattr(subprocess, 'CREATE_NO_WINDOW'):
            flags = subprocess.CREATE_NO_WINDOW
        result = subprocess.run(
            ["where", "tesseract"],
            capture_output=True, text=True, timeout=5,
            creationflags=flags)
        if result.returncode == 0:
            found = result.stdout.strip().split("\n")[0].strip()
            if found and os.path.isfile(found):
                candidates.insert(0, found)
    except Exception:
        pass

    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _configure_tesseract() -> Optional[str]:
    """Find Tesseract and configure pytesseract."""
    tess_path = shutil.which("tesseract")
    if tess_path is None and platform.system() == "Windows":
        tess_path = _find_tesseract_windows()
    if tess_path is None:
        return None

    try:
        result = subprocess.run(
            [tess_path, "--version"],
            capture_output=True, text=True, timeout=10)
        if result.returncode != 0:
            return None
    except Exception:
        return None

    try:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = tess_path
    except ImportError:
        return None
    return tess_path


# ═══════════════════════════════════════════════════════════════════════
#  BACKEND DETECTION
# ═══════════════════════════════════════════════════════════════════════
TESSERACT_PATH = _configure_tesseract()


def _tesseract_available() -> bool:
    if TESSERACT_PATH is None:
        return False
    try:
        import pytesseract
        return True
    except ImportError:
        return False


def _easyocr_available() -> bool:
    try:
        import easyocr
        return True
    except ImportError:
        return False


def _huggingface_available() -> bool:
    try:
        import transformers, torch
        return True
    except ImportError:
        return False


TESSERACT_OK = _tesseract_available()
EASYOCR_OK = _easyocr_available()
HUGGINGFACE_OK = _huggingface_available()
HAS_ANY_OCR = TESSERACT_OK or EASYOCR_OK or HUGGINGFACE_OK


def _default_engine() -> str:
    if TESSERACT_OK:
        return "pytesseract"
    if EASYOCR_OK:
        return "easyocr"
    if HUGGINGFACE_OK:
        return "huggingface"
    return "none"


DEFAULT_ENGINE = _default_engine()


# ═══════════════════════════════════════════════════════════════════════
#  LOGGING
# ═══════════════════════════════════════════════════════════════════════
LOG_FORMAT = "%(asctime)s | %(name)-28s | %(levelname)-7s | %(message)s"
LOG_FILE = LOG_DIR / "extraction.log"

_root_logger = logging.getLogger("soil_extractor")
if not _root_logger.handlers:
    _root_logger.setLevel(logging.INFO)
    _fmt = logging.Formatter(LOG_FORMAT)
    _fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    _fh.setFormatter(_fmt)
    _root_logger.addHandler(_fh)
    _sh = logging.StreamHandler()
    _sh.setFormatter(_fmt)
    _root_logger.addHandler(_sh)

logger = _root_logger


# ═══════════════════════════════════════════════════════════════════════
#  FILE LOADER
# ═══════════════════════════════════════════════════════════════════════
class FileLoader:

    @staticmethod
    def detect_type(path: str) -> Optional[str]:
        p = Path(path)
        if not p.is_file():
            return None
        ext = p.suffix.lower()
        if ext in SUPPORTED_IMAGE_EXT:
            return "image"
        if ext in SUPPORTED_PDF_EXT:
            return "pdf"
        if ext in SUPPORTED_WORD_EXT:
            return "word"
        return None

    @staticmethod
    def validate(path: str) -> Tuple[bool, str]:
        p = Path(path)
        if not p.exists():
            return False, f"File not found: {path}"
        if not p.is_file():
            return False, f"Not a file: {path}"
        if p.stat().st_size == 0:
            return False, f"Empty file: {path}"
        if p.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return False, f"File too large: {path}"
        if p.suffix.lower() not in SUPPORTED_EXT:
            return False, f"Unsupported format: {path}"
        return True, "ok"

    @staticmethod
    def find_files(directory: str) -> List[str]:
        d = Path(directory)
        if not d.is_dir():
            return []
        files = set()
        for ext in SUPPORTED_EXT:
            files.update(d.rglob(f"*{ext}"))
            files.update(d.rglob(f"*{ext.upper()}"))
        result = sorted(str(f.resolve()) for f in files)
        logger.info("Found %d file(s) in %s", len(result), directory)
        return result


# ═══════════════════════════════════════════════════════════════════════
#  IMAGE PROCESSOR
# ═══════════════════════════════════════════════════════════════════════
class ImageProcessor:

    @staticmethod
    def load(path: str) -> Optional[np.ndarray]:
        try:
            img = cv2.imread(path, cv2.IMREAD_COLOR)
            if img is None:
                pil = Image.open(path).convert("RGB")
                img = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)
            logger.info("Loaded image (%dx%d): %s",
                        img.shape[1], img.shape[0], path)
            return img
        except Exception as e:
            logger.error("Cannot load image %s: %s", path, e)
            return None

    @staticmethod
    def preprocess(image: np.ndarray) -> np.ndarray:
        try:
            proc = image.copy()
            if len(proc.shape) == 3:
                proc = cv2.cvtColor(proc, cv2.COLOR_BGR2GRAY)
            h, w = proc.shape[:2]
            if w < 1200:
                scale = 1200.0 / w
                proc = cv2.resize(proc, None, fx=scale, fy=scale,
                                  interpolation=cv2.INTER_CUBIC)
            proc = cv2.fastNlMeansDenoising(
                proc, None, DENOISE_STRENGTH, 7, 21)
            proc = cv2.adaptiveThreshold(
                proc, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, THRESHOLD_BLOCK_SIZE, THRESHOLD_C)
            proc = ImageProcessor._deskew(proc)
            return proc
        except Exception:
            if len(image.shape) == 3:
                return cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            return image

    @staticmethod
    def _deskew(img: np.ndarray, max_angle: float = 10.0) -> np.ndarray:
        try:
            coords = np.column_stack(np.where(img > 0))
            if len(coords) < 200:
                return img
            angle = cv2.minAreaRect(coords)[-1]
            angle = -(90 + angle) if angle < -45 else -angle
            if abs(angle) > max_angle or abs(angle) < 0.3:
                return img
            h, w = img.shape[:2]
            M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
            return cv2.warpAffine(img, M, (w, h),
                                  flags=cv2.INTER_CUBIC,
                                  borderMode=cv2.BORDER_REPLICATE)
        except Exception:
            return img


# ═══════════════════════════════════════════════════════════════════════
#  PDF PROCESSOR
# ═══════════════════════════════════════════════════════════════════════
class PDFProcessor:

    @staticmethod
    def to_images(pdf_path: str,
                  dpi: int = OCR_DPI) -> List[Image.Image]:
        try:
            from pdf2image import convert_from_path
            imgs = convert_from_path(pdf_path, dpi=dpi)
            logger.info("pdf2image: %d page(s)", len(imgs))
            return imgs
        except ImportError:
            pass
        except Exception as e:
            logger.warning("pdf2image failed: %s", e)

        try:
            import fitz
            doc = fitz.open(pdf_path)
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            images = []
            for page in doc:
                pix = page.get_pixmap(matrix=mat)
                images.append(Image.frombytes(
                    "RGB", (pix.width, pix.height), pix.samples))
            doc.close()
            return images
        except ImportError:
            logger.error("No PDF library installed")
            return []
        except Exception as e:
            logger.error("PyMuPDF failed: %s", e)
            return []

    @staticmethod
    def extract_text(pdf_path: str) -> Optional[str]:
        try:
            import fitz
            doc = fitz.open(pdf_path)
            parts = [page.get_text() for page in doc]
            doc.close()
            full = "\n".join(parts).strip()
            return full if len(full) > 80 else None
        except Exception:
            return None


# ═══════════════════════════════════════════════════════════════════════
#  WORD PROCESSOR
# ═══════════════════════════════════════════════════════════════════════
class WordProcessor:

    @staticmethod
    def extract_text(docx_path: str) -> Optional[str]:
        try:
            from docx import Document
            doc = Document(docx_path)
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells
                             if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
            return text if text else None
        except ImportError:
            logger.error("python-docx not installed")
            return None
        except Exception as e:
            logger.error("Word extraction failed: %s", e)
            return None

    @staticmethod
    def extract_images(docx_path: str) -> List[Image.Image]:
        try:
            from docx import Document
            doc = Document(docx_path)
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        images.append(
                            Image.open(io.BytesIO(
                                rel.target_part.blob)))
                    except Exception:
                        pass
            return images
        except Exception:
            return []


# ═══════════════════════════════════════════════════════════════════════
#  OCR RESULT
# ═══════════════════════════════════════════════════════════════════════
@dataclass
class OCRResult:
    text: str = ""
    confidence: float = 0.0
    word_confidences: List[Tuple[str, float]] = field(
        default_factory=list)
    engine_used: str = "unknown"


# ═══════════════════════════════════════════════════════════════════════
#  OCR ENGINE — multi-backend with smart merging
# ═══════════════════════════════════════════════════════════════════════
class OCREngine:

    def __init__(self, engine: str = None):
        requested = engine or DEFAULT_ENGINE
        self.engine = self._resolve(requested)
        self._easyocr_reader = None
        self._hf_proc = None
        self._hf_model = None
        self._img = ImageProcessor()

        if self.engine == "none":
            logger.error("NO OCR BACKEND AVAILABLE!")
        else:
            logger.info("OCR engine: requested=%s resolved=%s",
                        requested, self.engine)
            if self.engine in ("pytesseract", "both") and TESSERACT_PATH:
                logger.info("Tesseract: %s", TESSERACT_PATH)

    @staticmethod
    def _resolve(requested: str) -> str:
        avail = {
            "pytesseract": TESSERACT_OK,
            "easyocr": EASYOCR_OK,
            "huggingface": HUGGINGFACE_OK,
        }
        if requested == "both":
            if TESSERACT_OK and EASYOCR_OK:
                return "both"
            if TESSERACT_OK:
                return "pytesseract"
            if EASYOCR_OK:
                return "easyocr"
        if avail.get(requested, False):
            return requested
        for eng in ("pytesseract", "easyocr", "huggingface"):
            if avail.get(eng):
                logger.warning("'%s' unavailable → '%s'",
                               requested, eng)
                return eng
        return "none"

    def extract(self, image: Union[np.ndarray, Image.Image],
                preprocess: bool = True) -> OCRResult:
        """Run OCR. In 'both' mode, merges text from both engines."""
        if self.engine == "none":
            return OCRResult(engine_used="none")

        arr = np.array(image) if isinstance(
            image, Image.Image) else image
        processed = self._img.preprocess(arr) if preprocess else arr

        if self.engine == "both":
            return self._smart_both(processed, arr)

        dispatch = {
            "pytesseract": lambda: self._tess(processed),
            "easyocr": lambda: self._easy(arr),
            "huggingface": lambda: self._hf(arr),
        }

        result = dispatch.get(self.engine, lambda: OCRResult())()

        if not result.text.strip():
            result = self._fallback(arr, processed,
                                    exclude=self.engine)
        return result

    def _smart_both(self, processed: np.ndarray,
                    original: np.ndarray) -> OCRResult:
        """
        Run BOTH engines and combine their text output.

        Tesseract preserves line structure better.
        EasyOCR catches text Tesseract misses.
        We use Tesseract as primary and fill gaps with EasyOCR.
        """
        t_result = self._tess(processed)
        e_result = self._easy(original)

        t_text = t_result.text.strip()
        e_text = e_result.text.strip()

        logger.info(
            "Both engines: tesseract=%d chars (%.2f), "
            "easyocr=%d chars (%.2f)",
            len(t_text), t_result.confidence,
            len(e_text), e_result.confidence)

        # Strategy: Tesseract usually has better line structure.
        # If Tesseract got substantial text, prefer it.
        # Append EasyOCR text as supplementary at the end so
        # the parser has both versions to search through.

        combined_parts = []

        if t_text and len(t_text) > 50:
            combined_parts.append(t_text)

        if e_text and len(e_text) > 50:
            # Add EasyOCR text as a second block the parser can
            # also search through
            combined_parts.append(
                "--- EASYOCR_TEXT ---\n" + e_text)

        if not combined_parts:
            # Both failed; return whichever has more text
            if len(e_text) > len(t_text):
                return e_result
            return t_result

        combined = "\n\n".join(combined_parts)

        avg_conf = 0.0
        count = 0
        if t_result.confidence > 0:
            avg_conf += t_result.confidence
            count += 1
        if e_result.confidence > 0:
            avg_conf += e_result.confidence
            count += 1
        avg_conf = avg_conf / count if count else 0.0

        return OCRResult(
            text=combined,
            confidence=avg_conf,
            word_confidences=(t_result.word_confidences +
                              e_result.word_confidences),
            engine_used="both_merged",
        )

    def _fallback(self, orig, proc, exclude):
        attempts = []
        if exclude != "easyocr" and EASYOCR_OK:
            attempts.append(("easyocr", lambda: self._easy(orig)))
        if exclude != "pytesseract" and TESSERACT_OK:
            attempts.append(("pytesseract", lambda: self._tess(proc)))
        if exclude != "huggingface" and HUGGINGFACE_OK:
            attempts.append(("huggingface", lambda: self._hf(orig)))

        for name, fn in attempts:
            logger.info("Trying fallback: %s", name)
            try:
                r = fn()
                if r.text.strip():
                    r.engine_used = f"fallback_{name}"
                    return r
            except Exception:
                pass
        return OCRResult(engine_used="all_failed")

    def _tess(self, image: np.ndarray) -> OCRResult:
        try:
            import pytesseract
            if TESSERACT_PATH:
                pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

            pil = Image.fromarray(image)
            text = pytesseract.image_to_string(
                pil, config=TESSERACT_CONFIG)

            wc, confs = [], []
            try:
                d = pytesseract.image_to_data(
                    pil, output_type=pytesseract.Output.DICT,
                    config=TESSERACT_CONFIG)
                for j in range(len(d["text"])):
                    w = d["text"][j].strip()
                    c = int(d["conf"][j])
                    if w and c > 0:
                        wc.append((w, c / 100.0))
                        confs.append(c)
            except Exception:
                pass

            avg = (sum(confs) / len(confs) / 100.0) if confs else 0.0
            return OCRResult(text, avg, wc, "pytesseract")
        except Exception as e:
            logger.error("Tesseract error: %s", e)
            return OCRResult(engine_used="pytesseract_error")

    def _easy(self, image: np.ndarray) -> OCRResult:
        try:
            import easyocr
            if self._easyocr_reader is None:
                logger.info("Initialising EasyOCR…")
                self._easyocr_reader = easyocr.Reader(
                    EASYOCR_LANGUAGES, gpu=False, verbose=False)
                logger.info("EasyOCR ready")

            results = self._easyocr_reader.readtext(image)

            # Sort by vertical position then horizontal
            # to reconstruct reading order
            results.sort(key=lambda r: (
                min(pt[1] for pt in r[0]),
                min(pt[0] for pt in r[0])
            ))

            words, wc, confs = [], [], []
            prev_y = -999

            for bbox, txt, conf in results:
                curr_y = min(pt[1] for pt in bbox)
                # Insert newline when Y position jumps
                # (new line in the document)
                if prev_y >= 0 and (curr_y - prev_y) > 15:
                    words.append("\n")
                prev_y = curr_y

                words.append(txt)
                wc.append((txt, conf))
                confs.append(conf)

            text = " ".join(words)
            # Clean up newline spacing
            text = re.sub(r" ?\n ?", "\n", text)
            text = re.sub(r"\n{3,}", "\n\n", text)

            avg = (sum(confs) / len(confs)) if confs else 0.0
            return OCRResult(text, avg, wc, "easyocr")
        except Exception as e:
            logger.error("EasyOCR error: %s", e)
            return OCRResult(engine_used="easyocr_error")

    def _hf(self, image: np.ndarray) -> OCRResult:
        try:
            from transformers import (
                TrOCRProcessor, VisionEncoderDecoderModel)
            if self._hf_proc is None:
                name = "microsoft/trocr-base-printed"
                self._hf_proc = TrOCRProcessor.from_pretrained(name)
                self._hf_model = (
                    VisionEncoderDecoderModel.from_pretrained(name))
            pil = Image.fromarray(image).convert("RGB")
            px = self._hf_proc(
                images=pil, return_tensors="pt").pixel_values
            ids = self._hf_model.generate(px)
            text = self._hf_proc.batch_decode(
                ids, skip_special_tokens=True)[0]
            return OCRResult(text, 0.70,
                             engine_used="huggingface_trocr")
        except Exception as e:
            logger.error("TrOCR error: %s", e)
            return OCRResult(engine_used="huggingface_error")


# ═══════════════════════════════════════════════════════════════════════
#  PARSER — Completely rewritten with keyword-proximity extraction
# ═══════════════════════════════════════════════════════════════════════
@dataclass
class FieldResult:
    value: str = ""
    confidence: float = 0.0
    source: str = ""


@dataclass
class SoilReportData:
    farmer_details: Dict[str, str] = field(default_factory=lambda: {
        "name": "", "village": "", "district": "", "state": ""})
    sample_details: Dict[str, str] = field(default_factory=lambda: {
        "sample_id": "", "survey_number": "", "sample_date": ""})
    soil_parameters: Dict[str, str] = field(default_factory=lambda: {
        "ph": "", "electrical_conductivity": "",
        "organic_carbon": "",
        "nitrogen": "", "phosphorus": "", "potassium": "",
        "sulphur": "", "zinc": "", "iron": "",
        "copper": "", "manganese": "", "boron": ""})
    recommendations: Dict[str, str] = field(default_factory=lambda: {
        "fertilizer": "", "crop": ""})
    metadata: Dict[str, Any] = field(default_factory=dict)
    field_confidences: Dict[str, float] = field(
        default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "farmer_details": self.farmer_details,
            "sample_details": self.sample_details,
            "soil_parameters": self.soil_parameters,
            "recommendations": self.recommendations,
            "metadata": self.metadata,
            "field_confidences": self.field_confidences,
        }


class SoilReportParser:
    """
    Two-phase parser:
      Phase 1: Regex with lookahead boundaries
      Phase 2: Keyword-proximity fallback
    """

    # ── All known keywords (used as boundaries) ────────────────────
    ALL_KEYWORDS = [
        "farmer", "name", "village", "gram", "gaon",
        "district", "dist", "state",
        "sample", "survey", "khasra", "plot", "date",
        "ph", "electrical", "conductivity", "organic", "carbon",
        "nitrogen", "phosphorus", "potassium",
        "sulphur", "sulfur", "zinc", "iron",
        "copper", "manganese", "boron",
        "fertilizer", "crop", "recommendation",
        "---", "information", "results", "test",
    ]

    # Build lookahead boundary pattern
    _KW_BOUNDARY = (
        r"(?=\s*(?:"
        + "|".join([
            r"Village", r"Gram", r"Gaon",
            r"District", r"Dist",
            r"State", r"Rajya",
            r"Sample", r"Survey", r"Khasra", r"Plot",
            r"Date",
            r"pH", r"Electrical", r"Organic",
            r"Nitrogen", r"Phosphorus", r"Potassium",
            r"Sulphur", r"Sulfur", r"Zinc", r"Iron",
            r"Copper", r"Manganese", r"Boron",
            r"Fertilizer", r"Crop",
            r"Recommendation",
            r"---", r"\n\n",
        ])
        + r"|\Z))"
    )

    # Numeric boundary: stop before next word (not a unit)
    _NUM_STOP = (
        r"(?=\s*(?:dS|ds|ppm|mg|kg|%|"
        r"Electrical|Organic|Nitrogen|Phosphorus|Potassium|"
        r"Sulphur|Sulfur|Zinc|Iron|Copper|Manganese|Boron|"
        r"Fertilizer|Crop|Recommendation|"
        r"---|\n|\Z))"
    )

    # ── Valid numeric ranges ───────────────────────────────────────
    RANGES: Dict[str, Tuple[float, float]] = {
        "ph": (0, 14), "electrical_conductivity": (0, 30),
        "organic_carbon": (0, 15), "nitrogen": (0, 2000),
        "phosphorus": (0, 600), "potassium": (0, 2500),
        "sulphur": (0, 500), "zinc": (0, 100),
        "iron": (0, 500), "copper": (0, 100),
        "manganese": (0, 500), "boron": (0, 50),
    }

    # ── Keyword → extraction config map ────────────────────────────
    # Each entry: (keyword_patterns, capture_type, post_keywords)
    # capture_type: "text", "numeric", "date", "block"

    def __init__(self):
        self._text = ""
        self._ocr_conf = 0.0
        self._results: Dict[str, FieldResult] = {}

    def parse(self, text: str,
              ocr_confidence: float = 0.0) -> SoilReportData:
        """Parse OCR text into structured SoilReportData."""
        self._text = self._clean(text)
        self._ocr_conf = ocr_confidence
        self._results = {}

        logger.debug("Cleaned text for parsing:\n%s",
                     self._text[:500])

        rpt = SoilReportData()
        rpt.farmer_details = self._farmer()
        rpt.sample_details = self._sample()
        rpt.soil_parameters = self._params()
        rpt.recommendations = self._recs()
        rpt.field_confidences = self._confidences(rpt)
        rpt.metadata = self._meta(rpt)
        return rpt

    # ── Text cleaning ──────────────────────────────────────────────
    def _clean(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"\s*\|\s*", " : ", text)

        # Fix common OCR errors
        replacements = [
            (r"\bph\b", "pH"),
            (r"\bPh\b", "pH"),
            (r"\bPH\b", "pH"),
            (r"\bkg/haa?\b", "kg/ha"),
            (r"\bds/m\b", "dS/m"),
            (r"\bkgha\b", "kg/ha"),
            (r"\bkg ha\b", "kg/ha"),
        ]
        for pat, repl in replacements:
            text = re.sub(pat, repl, text)

        # Strip non-ASCII
        text = re.sub(r"[^\x20-\x7E\n\r\t]", " ", text)
        return text.strip()

    # ── Core extraction: keyword-proximity method ──────────────────
    def _extract_after_keyword(
        self,
        keywords: List[str],
        stop_keywords: List[str],
        capture: str = "text",
        label: str = "",
    ) -> str:
        """
        Find one of `keywords` in text, then capture the value
        that follows, stopping at `stop_keywords` or next line.

        capture: "text" | "numeric" | "date" | "block"
        """
        text = self._text

        for kw in keywords:
            # Build pattern: keyword followed by optional separator
            # then capture group
            escaped = re.escape(kw).replace(r"\ ", r"\s+")
            pat = (
                rf"(?:{escaped})"
                r"\s*(?:\([^)]*\))?"   # optional parenthetical
                r"\s*[:\-=|]*\s*"       # optional separator
            )

            m = re.search(pat, text, re.IGNORECASE)
            if not m:
                continue

            # Text after the keyword match
            after = text[m.end():]

            if capture == "numeric":
                val = self._grab_number(after)
            elif capture == "date":
                val = self._grab_date(after)
            elif capture == "block":
                val = self._grab_block(after, stop_keywords)
            else:  # "text"
                val = self._grab_text(after, stop_keywords)

            if val:
                conf = max(0.6, self._ocr_conf) if self._ocr_conf > 0 else 0.5
                self._results[label] = FieldResult(val, conf, kw)
                return val

        return ""

    def _grab_text(self, after: str,
                   stop_kws: List[str]) -> str:
        """Grab text value, stopping at next known keyword."""
        # Build stop pattern
        stop_parts = [re.escape(k) for k in stop_kws]
        stop_parts.extend([r"---", r"\n\n"])
        stop_re = "|".join(stop_parts)

        # Grab up to 60 chars or until stop keyword
        m = re.match(
            rf"([A-Za-z\s\.\,\-]{{2,60}}?)(?=\s*(?:{stop_re})|\s*$)",
            after,
            re.IGNORECASE
        )
        if m:
            return m.group(1).strip()

        # Fallback: grab until newline or stop word
        m = re.match(
            r"([A-Za-z\s\.\,\-]{2,60})",
            after.split("\n")[0],
        )
        if m:
            val = m.group(1).strip()
            # Post-trim: remove trailing stop keywords
            for sk in stop_kws:
                idx = val.lower().find(sk.lower())
                if idx > 0:
                    val = val[:idx].strip()
            return val

        return ""

    def _grab_number(self, after: str) -> str:
        """Grab first numeric value after position."""
        m = re.match(r"\s*(\d+\.?\d*)", after)
        return m.group(1) if m else ""

    def _grab_date(self, after: str) -> str:
        """Grab a date pattern."""
        m = re.match(
            r"\s*(\d{1,2}[\-/\.]\d{1,2}[\-/\.]\d{2,4})",
            after)
        return m.group(1).strip() if m else ""

    def _grab_block(self, after: str,
                    stop_kws: List[str]) -> str:
        """Grab a longer text block (for recommendations)."""
        stop_parts = [re.escape(k) for k in stop_kws]
        stop_parts.extend([r"\n\n", r"---"])
        stop_re = "|".join(stop_parts)

        m = re.match(
            rf"([\s\S]{{10,400}}?)(?=\s*(?:{stop_re})|\s*$)",
            after,
            re.IGNORECASE
        )
        if m:
            return re.sub(r"\s+", " ", m.group(1)).strip()

        # Fallback: take up to 300 chars
        val = after[:300].strip()
        for sk in stop_kws:
            idx = val.lower().find(sk.lower())
            if idx > 0:
                val = val[:idx].strip()
        return re.sub(r"\s+", " ", val).strip(" :-=") if val else ""

    # ── Farmer Details ─────────────────────────────────────────────
    def _farmer(self) -> dict:
        name = self._extract_after_keyword(
            keywords=[
                "Farmer Name", "Farmer's Name",
                "farmer name", "Name of Farmer",
                "Name of the Farmer",
                "Applicant Name", "Beneficiary Name",
                "Kisan Name",
            ],
            stop_keywords=[
                "Village", "Gram", "Gaon", "District",
                "State", "Sample", "Survey", "---",
            ],
            capture="text",
            label="farmer_name",
        )

        village = self._extract_after_keyword(
            keywords=[
                "Village", "Gram", "Gaon", "Grama",
                "Habitation", "Town",
            ],
            stop_keywords=[
                "District", "Dist", "State", "Sample",
                "Survey", "---", "Farmer",
            ],
            capture="text",
            label="village",
        )

        district = self._extract_after_keyword(
            keywords=[
                "District", "Dist", "Jila", "Zilla",
            ],
            stop_keywords=[
                "State", "Sample", "Survey", "Village",
                "---", "Farmer",
            ],
            capture="text",
            label="district",
        )

        state = self._extract_after_keyword(
            keywords=[
                "State", "Rajya",
            ],
            stop_keywords=[
                "Sample", "Survey", "District", "Village",
                "Farmer", "---", "Information",
            ],
            capture="text",
            label="state",
        )

        return {
            "name": self._title_case(name),
            "village": self._title_case(village),
            "district": self._title_case(district),
            "state": self._title_case(state),
        }

    # ── Sample Details ─────────────────────────────────────────────
    def _sample(self) -> dict:
        sample_id = self._extract_after_keyword(
            keywords=[
                "Sample ID", "Sample No", "Sample Number",
                "Sample Code", "SHC ID", "SHC No",
                "Lab Sample No", "Lab No",
            ],
            stop_keywords=[
                "Survey", "Sample Date", "Date", "pH",
                "---", "Soil",
            ],
            capture="text",
            label="sample_id",
        )
        # For sample ID, also try to capture alphanumeric codes
        if not sample_id:
            for kw in ["Sample ID", "SHC"]:
                pat = (
                    rf"{kw}\s*[:\-=|]*\s*"
                    r"([A-Za-z0-9][\w\-/]+)"
                )
                m = re.search(pat, self._text, re.IGNORECASE)
                if m:
                    sample_id = m.group(1).strip()
                    break

        survey = self._extract_after_keyword(
            keywords=[
                "Survey Number", "Survey No",
                "Khasra Number", "Khasra No",
                "Plot Number", "Plot No",
            ],
            stop_keywords=[
                "Sample Date", "Date", "pH", "---",
                "Soil", "Farmer",
            ],
            capture="text",
            label="survey_no",
        )
        # Grab alphanumeric survey number
        if not survey:
            for kw in ["Survey Number", "Survey No",
                        "Khasra", "Plot"]:
                pat = (
                    rf"{kw}\s*[:\-=|]*\s*"
                    r"([A-Za-z0-9][\w\-/]*)"
                )
                m = re.search(pat, self._text, re.IGNORECASE)
                if m:
                    survey = m.group(1).strip()
                    break

        date_str = self._extract_after_keyword(
            keywords=[
                "Sample Date", "Sample Collection Date",
                "Date of Collection", "Collection Date",
                "Date of Sample",
            ],
            stop_keywords=[
                "pH", "Soil", "---", "Electrical",
                "Results",
            ],
            capture="date",
            label="sample_date",
        )
        # Broader date fallback
        if not date_str:
            m = re.search(
                r"(?:Date)\s*[:\-=|]*\s*"
                r"(\d{1,2}[\-/\.]\d{1,2}[\-/\.]\d{2,4})",
                self._text, re.IGNORECASE)
            if m:
                date_str = m.group(1)

        return {
            "sample_id": sample_id,
            "survey_number": survey,
            "sample_date": self._norm_date(date_str),
        }

    # ── Soil Parameters ────────────────────────────────────────────
    def _params(self) -> dict:
        param_config = {
            "ph": [
                "pH", "Soil pH", "p.H",
            ],
            "electrical_conductivity": [
                "Electrical Conductivity (EC)",
                "Electrical Conductivity",
                "EC", "E.C.", "E.C",
                "Elec. Conductivity",
            ],
            "organic_carbon": [
                "Organic Carbon (OC)",
                "Organic Carbon",
                "OC", "O.C.", "O.C",
                "Org. Carbon", "Org Carbon",
            ],
            "nitrogen": [
                "Nitrogen (N)", "Nitrogen",
                "Available N", "N (kg",
                "Available Nitrogen",
            ],
            "phosphorus": [
                "Phosphorus (P)", "Phosphorus",
                "Available P", "P2O5", "P (kg",
                "Available Phosphorus",
            ],
            "potassium": [
                "Potassium (K)", "Potassium",
                "Available K", "K2O", "K (kg",
                "Available Potassium",
            ],
            "sulphur": [
                "Sulphur (S)", "Sulphur", "Sulfur (S)",
                "Sulfur", "Available S",
            ],
            "zinc": [
                "Zinc (Zn)", "Zinc",
                "Available Zn", "Zn",
            ],
            "iron": [
                "Iron (Fe)", "Iron",
                "Available Fe", "Fe",
            ],
            "copper": [
                "Copper (Cu)", "Copper",
                "Available Cu", "Cu",
            ],
            "manganese": [
                "Manganese (Mn)", "Manganese",
                "Available Mn", "Mn",
            ],
            "boron": [
                "Boron (B)", "Boron",
                "Available B",
            ],
        }

        # Build the stop keywords list for numeric params:
        # next parameter names + common non-numeric words
        all_param_names = [
            "Electrical", "Organic", "Nitrogen", "Phosphorus",
            "Potassium", "Sulphur", "Sulfur", "Zinc", "Iron",
            "Copper", "Manganese", "Boron",
            "Fertilizer", "Crop", "Recommendation", "---",
        ]

        results = {}
        for param_name, keywords in param_config.items():
            val = self._extract_after_keyword(
                keywords=keywords,
                stop_keywords=all_param_names,
                capture="numeric",
                label=param_name,
            )
            val = self._validate_number(val, param_name)
            results[param_name] = val

        return results

    # ── Recommendations ────────────────────────────────────────────
    def _recs(self) -> dict:
        fert = self._extract_after_keyword(
            keywords=[
                "Fertilizer Recommendation",
                "Fertiliser Recommendation",
                "Recommended Fertilizer",
                "Recommended Fertiliser",
                "Fertilizer Dose",
                "Fertiliser Dose",
            ],
            stop_keywords=[
                "Crop Recommendation", "Crop Suggested",
                "Suitable Crop", "Recommended Crop",
                "---",
            ],
            capture="block",
            label="fertilizer",
        )

        crop = self._extract_after_keyword(
            keywords=[
                "Crop Recommendation",
                "Recommended Crops",
                "Suitable Crops",
                "Crop Suggested",
                "Recommended Crop",
                "Suitable Crop",
            ],
            stop_keywords=[
                "Fertilizer", "Fertiliser", "---",
                "Note", "Remark",
            ],
            capture="block",
            label="crop",
        )

        # Clean recommendations
        fert = self._clean_rec(fert)
        crop = self._clean_rec(crop)

        return {
            "fertilizer": fert,
            "crop": crop,
        }

    # ── Value cleaners ─────────────────────────────────────────────
    @staticmethod
    def _title_case(s: str) -> str:
        if not s:
            return ""
        s = re.sub(r"[:\-=]+$", "", s).strip()
        s = re.sub(r"^\s*[:\-=]+", "", s).strip()
        # Remove leftover section headers
        for noise in ["Information", "Details", "---",
                       "Farmer", "Section"]:
            s = re.sub(
                rf"\b{noise}\b", "", s, flags=re.IGNORECASE
            ).strip()
        s = re.sub(r"\s+", " ", s).strip()
        if len(s) < 2 or len(s) > 60:
            return ""
        return s.title()

    @staticmethod
    def _norm_date(s: str) -> str:
        if not s:
            return ""
        s = s.strip()
        for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y",
                     "%d/%m/%y", "%d-%m-%y", "%m/%d/%Y",
                     "%Y-%m-%d"):
            try:
                return datetime.strptime(s, fmt).strftime(
                    "%d/%m/%Y")
            except ValueError:
                continue
        return s

    def _validate_number(self, val: str, param: str) -> str:
        if not val:
            return ""
        try:
            num = float(val)
        except ValueError:
            return ""
        lo, hi = self.RANGES.get(param, (0, 99999))
        if not (lo <= num <= hi):
            logger.warning("%s=%.2f outside [%.0f,%.0f]",
                           param, num, lo, hi)
            return ""
        return f"{num:g}"

    @staticmethod
    def _clean_rec(s: str) -> str:
        if not s:
            return ""
        s = re.sub(r"\s+", " ", s).strip(" :-=")
        # Remove trailing noise words
        for noise in ["Village", "Farmer", "Sample",
                       "District", "State", "---"]:
            idx = s.rfind(noise)
            if idx > 10:
                s = s[:idx].strip(" :-=,;")
        return s

    # ── Confidence / metadata ──────────────────────────────────────
    def _confidences(self, rpt: SoilReportData) -> dict:
        confs = {}
        for section, prefix in [
            (rpt.farmer_details, "farmer"),
            (rpt.sample_details, "sample"),
            (rpt.soil_parameters, "param"),
            (rpt.recommendations, "rec"),
        ]:
            for key, val in section.items():
                label = f"{prefix}_{key}"
                fr = self._results.get(key) or \
                     self._results.get(label)
                if val and fr:
                    confs[label] = round(fr.confidence, 3)
                elif val:
                    confs[label] = round(
                        self._ocr_conf * 0.8, 3)
                else:
                    confs[label] = 0.0
        return confs

    def _meta(self, rpt: SoilReportData) -> dict:
        all_v = (
            list(rpt.farmer_details.values()) +
            list(rpt.sample_details.values()) +
            list(rpt.soil_parameters.values()) +
            list(rpt.recommendations.values()))
        filled = sum(bool(v) for v in all_v)
        total = len(all_v)
        return {
            "extraction_confidence": round(
                filled / total, 3) if total else 0.0,
            "fields_extracted": filled,
            "total_fields": total,
            "ocr_engine": "",
            "processing_date": datetime.now().isoformat(),
        }


# ═══════════════════════════════════════════════════════════════════════
#  JSON WRITER
# ═══════════════════════════════════════════════════════════════════════
class JSONWriter:

    @staticmethod
    def write(data: dict, output_path: str = None,
              source_file: str = None) -> str:
        if output_path is None:
            stem = (Path(source_file).stem if source_file
                    else datetime.now().strftime("%Y%m%d_%H%M%S"))
            output_path = str(
                OUTPUT_DIR / f"{stem}_extracted.json")
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w", encoding="utf-8") as fp:
            json.dump(data, fp, indent=2, ensure_ascii=False)
        logger.info("JSON → %s", out)
        return str(out.resolve())

    @staticmethod
    def write_batch(results: List[dict],
                    output_path: str = None) -> str:
        if output_path is None:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(
                OUTPUT_DIR / f"batch_{ts}.json")
        ok = sum(1 for r in results
                 if r.get("status") == "success")
        batch = {
            "batch_info": {
                "total_files": len(results),
                "successful": ok,
                "failed": len(results) - ok,
                "date": datetime.now().isoformat(),
            },
            "results": results,
        }
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w", encoding="utf-8") as fp:
            json.dump(batch, fp, indent=2, ensure_ascii=False)
        logger.info("Batch report → %s", out)
        return str(out.resolve())


# ═══════════════════════════════════════════════════════════════════════
#  SAMPLE REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════════════
def create_sample_report() -> str:
    W, H = 900, 1300
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    try:
        ft = ImageFont.truetype("arial.ttf", 22)
        fh = ImageFont.truetype("arial.ttf", 16)
        fb = ImageFont.truetype("arial.ttf", 14)
    except OSError:
        ft = fh = fb = ImageFont.load_default()

    y = 30

    def line(text, font=fb, indent=40, gap=26):
        nonlocal y
        draw.text((indent, y), text, fill="black", font=font)
        y += gap

    draw.text((200, y), "SOIL HEALTH CARD REPORT",
              fill="darkgreen", font=ft)
    y += 50

    line("--- Farmer Information ---", fh, 30, 30)
    line("Farmer Name  : Ramesh Kumar Sharma")
    line("Village      : Sundarpur")
    line("District     : Varanasi")
    line("State        : Uttar Pradesh")
    y += 10

    line("--- Sample Information ---", fh, 30, 30)
    line("Sample ID          : SHC-2024-UP-00456")
    line("Survey Number      : 123/4A")
    line("Sample Date        : 15/03/2024")
    y += 10

    line("--- Soil Test Results ---", fh, 30, 30)
    for name, val, unit in [
        ("pH", "7.5", ""),
        ("Electrical Conductivity (EC)", "0.45", "dS/m"),
        ("Organic Carbon (OC)", "0.62", "%"),
        ("Nitrogen (N)", "245.0", "kg/ha"),
        ("Phosphorus (P)", "18.5", "kg/ha"),
        ("Potassium (K)", "312.0", "kg/ha"),
        ("Sulphur (S)", "14.2", "ppm"),
        ("Zinc (Zn)", "0.85", "ppm"),
        ("Iron (Fe)", "6.20", "ppm"),
        ("Copper (Cu)", "1.10", "ppm"),
        ("Manganese (Mn)", "3.45", "ppm"),
        ("Boron (B)", "0.52", "ppm"),
    ]:
        suf = f" {unit}" if unit else ""
        line(f"  {name:38s}: {val}{suf}")
    y += 10

    line("--- Recommendations ---", fh, 30, 30)
    line("Fertilizer Recommendation: Apply Urea 120 kg/ha, "
         "DAP 60 kg/ha,")
    line("    MOP 40 kg/ha in split doses as basal and "
         "top dressing.")
    y += 4
    line("Crop Recommendation: Wheat, Rice, Mustard, "
         "Sugarcane")

    draw.rectangle([(10, 10), (W - 10, H - 10)],
                   outline="darkgreen", width=2)

    out = INPUT_DIR / "sample_soil_report.png"
    img.save(out)
    print(f"  Sample report → {out}")
    return str(out)


# ═══════════════════════════════════════════════════════════════════════
#  ENVIRONMENT CHECKER
# ═══════════════════════════════════════════════════════════════════════
def check_environment() -> bool:
    def _try(name):
        try:
            __import__(name)
            return True
        except ImportError:
            return False

    print("\n  ╔══════════════════════════════════════════════╗")
    print("  ║  Soil Report Extractor — Environment Check  ║")
    print("  ╚══════════════════════════════════════════════╝\n")

    tess_found = TESSERACT_PATH is not None
    checks = [
        ("Python " + ".".join(map(str, sys.version_info[:3])),
         sys.version_info >= (3, 8), "Install Python 3.8+"),
        ("Pillow", _try("PIL"), "pip install Pillow"),
        ("numpy", _try("numpy"), "pip install numpy"),
        ("opencv-python", _try("cv2"),
         "pip install opencv-python"),
        ("Tesseract binary", tess_found,
         "Install Tesseract"),
        ("pytesseract", _try("pytesseract"),
         "pip install pytesseract"),
        ("easyocr", EASYOCR_OK, "pip install easyocr"),
        ("PyMuPDF", _try("fitz"), "pip install PyMuPDF"),
        ("pdf2image", _try("pdf2image"),
         "pip install pdf2image"),
        ("python-docx", _try("docx"),
         "pip install python-docx"),
    ]
    for label, ok, fix in checks:
        icon = "✅" if ok else "❌"
        suffix = "" if ok else f"  ← {fix}"
        print(f"    {icon}  {label}{suffix}")

    if tess_found:
        print(f"\n    Tesseract: {TESSERACT_PATH}")

    print()
    if HAS_ANY_OCR:
        print(f"  ✅  READY — default engine: {DEFAULT_ENGINE}")
    else:
        print("  ❌  NO OCR BACKEND!")
        print("     pip install easyocr")
    print()
    return HAS_ANY_OCR


# ═══════════════════════════════════════════════════════════════════════
#  MAIN PIPELINE
# ═══════════════════════════════════════════════════════════════════════
class SoilReportExtractor:

    def __init__(self, engine: str = None):
        self.loader = FileLoader()
        self.pdf = PDFProcessor()
        self.img_proc = ImageProcessor()
        self.word = WordProcessor()
        self.ocr = OCREngine(engine=engine)
        self.parser = SoilReportParser()
        self.writer = JSONWriter()
        logger.info("Pipeline ready (engine=%s)",
                    self.ocr.engine)

    def process_file(self, file_path: str,
                     output_path: str = None) -> Dict[str, Any]:
        t0 = time.time()
        logger.info("▶ Processing: %s", file_path)

        ok, msg = self.loader.validate(file_path)
        if not ok:
            logger.error(msg)
            return {"source_file": file_path,
                    "status": "error",
                    "error": msg, "processing_time": 0}

        try:
            ftype = self.loader.detect_type(file_path)
            ocr_result = self._route(file_path, ftype)

            if not ocr_result.text.strip():
                logger.warning("No text from %s", file_path)
                return {
                    "source_file": file_path,
                    "status": "warning",
                    "warning": "No text extracted. "
                               "Run --check",
                    "data": self.parser.parse("").to_dict(),
                    "processing_time": round(
                        time.time() - t0, 2),
                }

            report = self.parser.parse(
                ocr_result.text, ocr_result.confidence)
            report.metadata["ocr_engine"] = \
                ocr_result.engine_used
            data = report.to_dict()

            jp = self.writer.write(
                data, output_path=output_path,
                source_file=file_path)
            elapsed = round(time.time() - t0, 2)
            logger.info("✔ %s in %.2fs", file_path, elapsed)

            return {
                "source_file": file_path,
                "status": "success",
                "output_file": jp,
                "data": data,
                "processing_time": elapsed,
            }

        except Exception as exc:
            logger.error("✘ %s: %s", file_path, exc,
                         exc_info=True)
            return {
                "source_file": file_path,
                "status": "error",
                "error": str(exc),
                "processing_time": round(
                    time.time() - t0, 2),
            }

    def _route(self, path: str, ftype: str) -> OCRResult:
        if ftype == "image":
            return self._do_image(path)
        if ftype == "pdf":
            return self._do_pdf(path)
        if ftype == "word":
            return self._do_word(path)
        raise ValueError(f"Unsupported: {ftype}")

    def _do_image(self, path: str) -> OCRResult:
        arr = self.img_proc.load(path)
        if arr is None:
            return OCRResult()
        return self.ocr.extract(arr)

    def _do_pdf(self, path: str) -> OCRResult:
        embedded = self.pdf.extract_text(path)
        if embedded and len(embedded.strip()) > 100:
            return OCRResult(embedded, 0.95,
                             engine_used="pdf_embedded")
        pages = self.pdf.to_images(path)
        if not pages:
            return OCRResult()
        texts, confs = [], []
        for i, pil in enumerate(pages):
            logger.info("  page %d/%d", i + 1, len(pages))
            r = self.ocr.extract(np.array(pil))
            texts.append(r.text)
            confs.append(r.confidence)
        return OCRResult(
            "\n\n".join(texts),
            sum(confs) / len(confs) if confs else 0.0,
            engine_used=f"pdf_ocr_{len(pages)}pg")

    def _do_word(self, path: str) -> OCRResult:
        text = self.word.extract_text(path)
        if text and len(text.strip()) > 60:
            return OCRResult(text, 0.95,
                             engine_used="docx_direct")
        images = self.word.extract_images(path)
        if not images:
            return OCRResult()
        texts, confs = [], []
        for pil in images:
            r = self.ocr.extract(np.array(pil))
            texts.append(r.text)
            confs.append(r.confidence)
        return OCRResult(
            "\n".join(texts),
            sum(confs) / len(confs) if confs else 0.0,
            engine_used="docx_image_ocr")

    def process_batch(self, paths: List[str],
                      workers: int = MAX_WORKERS) -> List[dict]:
        logger.info("Batch: %d file(s)", len(paths))
        results = []
        with ThreadPoolExecutor(max_workers=workers) as pool:
            futs = {pool.submit(self.process_file, p): p
                    for p in paths}
            for f in as_completed(futs):
                try:
                    results.append(f.result())
                except Exception as e:
                    results.append({
                        "source_file": futs[f],
                        "status": "error",
                        "error": str(e)})
        self.writer.write_batch(results)
        return results

    def process_directory(self, directory: str,
                          workers: int = MAX_WORKERS):
        paths = self.loader.find_files(directory)
        if not paths:
            logger.warning("No files in %s", directory)
            return []
        return self.process_batch(paths, workers)


# ═══════════════════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════════════════
def _print_result(r: dict) -> None:
    src = r.get("source_file", "?")
    st = r.get("status", "?")
    if st == "success":
        m = r.get("data", {}).get("metadata", {})
        print(f"\n{'=' * 64}")
        print(f"  ✅  {src}")
        print(f"      Output  : {r.get('output_file')}")
        print(f"      Time    : {r.get('processing_time')}s")
        print(f"      Fields  : "
              f"{m.get('fields_extracted', 0)}/"
              f"{m.get('total_fields', 0)}")
        print(f"      Conf    : "
              f"{m.get('extraction_confidence', 0):.1%}")
        print(f"{'=' * 64}")

        d = r.get("data", {})
        print("\n  Farmer Details:")
        for k, v in d.get("farmer_details", {}).items():
            print(f"    {k:12s}: {v or '(not found)'}")
        print("  Sample Details:")
        for k, v in d.get("sample_details", {}).items():
            print(f"    {k:14s}: {v or '(not found)'}")
        print("  Soil Parameters:")
        for k, v in d.get("soil_parameters", {}).items():
            print(f"    {k:25s}: {v or '(not found)'}")
        print("  Recommendations:")
        for k, v in d.get("recommendations", {}).items():
            print(f"    {k:12s}: {v or '(not found)'}")
        print()
    elif st == "warning":
        print(f"\n  ⚠️   {src} — {r.get('warning', '')}")
    else:
        print(f"\n  ❌  {src} — {r.get('error', 'unknown')}")


def main():
    ap = argparse.ArgumentParser(
        description="Soil Test Report OCR Extractor",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python soil_extractor.py --check
  python soil_extractor.py --create-sample
  python soil_extractor.py -i input/sample_soil_report.png
  python soil_extractor.py -i report.pdf -e easyocr
  python soil_extractor.py -i report.pdf -e both
  python soil_extractor.py -d ./input/ -w 4
""")
    ap.add_argument("-i", "--input", nargs="+",
                    help="Input file(s)")
    ap.add_argument("-d", "--input-dir",
                    help="Process directory")
    ap.add_argument("-o", "--output-dir",
                    default=str(OUTPUT_DIR))
    ap.add_argument("-e", "--engine",
                    choices=["pytesseract", "easyocr",
                             "both", "huggingface"],
                    default=None)
    ap.add_argument("-w", "--workers", type=int,
                    default=MAX_WORKERS)
    ap.add_argument("-v", "--verbose", action="store_true")
    ap.add_argument("--check", action="store_true")
    ap.add_argument("--create-sample", action="store_true")

    args = ap.parse_args()

    if args.check:
        check_environment()
        sys.exit(0)
    if args.create_sample:
        create_sample_report()
        sys.exit(0)
    if args.verbose:
        _root_logger.setLevel(logging.DEBUG)

    Path(args.output_dir).mkdir(parents=True, exist_ok=True)

    if not HAS_ANY_OCR:
        print()
        check_environment()
        sys.exit(1)

    ext = SoilReportExtractor(engine=args.engine)

    if args.input:
        if len(args.input) == 1:
            _print_result(ext.process_file(args.input[0]))
        else:
            for r in ext.process_batch(
                    args.input, args.workers):
                _print_result(r)
    elif args.input_dir:
        for r in ext.process_directory(
                args.input_dir, args.workers):
            _print_result(r)
    else:
        ap.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()