"""
Word Processor Module
=====================
Extracts text, tables, and images from .docx files.
"""

import io
from typing import List, Optional
import logging
from PIL import Image

logger = logging.getLogger("soil_report_extractor.word_processor")


class WordProcessor:
    """Extract content from .docx documents."""

    @staticmethod
    def extract_text(docx_path: str) -> Optional[str]:
        """Return all readable text (paragraphs + table cells)."""
        try:
            from docx import Document
            doc = Document(docx_path)
            parts = []

            for para in doc.paragraphs:
                line = para.text.strip()
                if line:
                    parts.append(line)

            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            text = "\n".join(parts)
            if text:
                logger.info("Extracted %d chars from %s", len(text), docx_path)
                return text
            return None
        except ImportError:
            logger.error("python-docx is not installed")
            return None
        except Exception as exc:
            logger.error("Word extraction failed for %s: %s", docx_path, exc)
            return None

    @staticmethod
    def extract_images_from_docx(docx_path: str) -> List[Image.Image]:
        """Extract embedded images as PIL Images."""
        try:
            from docx import Document
            doc = Document(docx_path)
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        blob = rel.target_part.blob
                        images.append(Image.open(io.BytesIO(blob)))
                    except Exception:
                        pass
            logger.info("Extracted %d image(s) from %s", len(images), docx_path)
            return images
        except Exception as exc:
            logger.error("Image extraction failed for %s: %s", docx_path, exc)
            return []